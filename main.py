from googletrans import Translator
import tkinter as tk
from tkinter import ttk, filedialog
import os
from docx import Document
from zipfile import ZipFile

class TradocsApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Tradocs")

        self.files = []
        self.translator = Translator()

        # Interface
        self.file_listbox = tk.Listbox(root, selectmode=tk.MULTIPLE)
        self.file_listbox.pack(pady=10)

        self.select_button = tk.Button(root, text="Sélectionner les fichiers", command=self.select_files)
        self.select_button.pack(pady=5)

        self.lang_from_label = tk.Label(root, text="Langue source:")
        self.lang_from_label.pack()

        self.lang_from_combobox = ttk.Combobox(root, values=self.get_supported_languages())
        self.lang_from_combobox.pack()

        self.lang_to_label = tk.Label(root, text="Langue cible:")
        self.lang_to_label.pack()

        self.lang_to_combobox = ttk.Combobox(root, values=self.get_supported_languages())
        self.lang_to_combobox.pack()

        self.translate_button = tk.Button(root, text="Traduire", command=self.translate_files)
        self.translate_button.pack(pady=10)

        self.download_button = tk.Button(root, text="Télécharger les fichiers traduits", command=self.download_files)
        self.download_button.pack()

        self.file_status_label = tk.Label(root, text="Statut des fichiers:")
        self.file_status_label.pack()

        self.file_status_listbox = tk.Listbox(root)
        self.file_status_listbox.pack()

    def get_supported_languages(self):
        return ["en", "fr", "es", "de", "it", "ja", "ko", "zh", "ru", "ar"]  # Ajout de l'anglais dans la liste

    def select_files(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("Word Files", "*.docx")])
        self.files.extend(file_paths)

        for file_path in file_paths:
            self.file_listbox.insert(tk.END, os.path.basename(file_path))
            self.file_status_listbox.insert(tk.END, "Non traduit")

    def translate_files(self):
        lang_from = self.lang_from_combobox.get()
        lang_to = self.lang_to_combobox.get()

        for i, file_path in enumerate(self.files):
            doc = Document(file_path)
        
            try:
                translated_text = self.translator.translate("\n".join([p.text for p in doc.paragraphs]), src=lang_from, dest=lang_to).text
            except Exception as e:
                print(f"Erreur lors de la traduction du fichier {file_path}: {e}")
                translated_text = "Erreur de traduction"

            # Supprimer le contenu du document en supprimant tous les paragraphes existants
            for p in doc.paragraphs:
                p.clear()

            # Ajouter les paragraphes traduits
            for para in translated_text.split('\n'):
                doc.add_paragraph(para)

            translated_file_path = file_path.replace(".docx", f"_translated_{lang_to}.docx")
            doc.save(translated_file_path)

            self.file_status_listbox.delete(i)
            self.file_status_listbox.insert(i, "Traduit")



    def download_files(self):
        zip_filename = "translated_files.zip"
        with ZipFile(zip_filename, 'w') as zipf:
            for file_path in self.files:
                translated_file_path = file_path.replace(".docx", f"_translated_{self.lang_to_combobox.get()}.docx")
                zipf.write(translated_file_path, os.path.basename(translated_file_path))

        os.remove(translated_file_path)
        self.root.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    app = TradocsApp(root)
    root.mainloop()
